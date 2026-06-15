import io
import re
from datetime import datetime
from urllib.parse import quote
from typing import List, Optional, Dict, Any

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

# Пытаемся импортировать библиотеки для парсинга документов
try:
    import docx
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None

router = APIRouter()


class DocumentLine(BaseModel):
    text: str
    font: str
    size: float
    page: int


# Временное хранилище для MVP (демонстрации)
audit_history = []


def extract_lines_from_file(file_name: str, file_bytes: bytes) -> List[DocumentLine]:
    """Извлекает текст, шрифты и размеры строк из загруженного файла"""
    lines = []
    file_ext = file_name.split('.')[-1].lower()

    # 1. Парсинг файлов Microsoft Word (.docx)
    if file_ext == 'docx' and docx:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                # Извлекаем шрифт и размер из первого фрагмента текста (run)
                font_name = "Times New Roman"
                font_size = 14.0
                if paragraph.runs:
                    run = paragraph.runs[0]
                    if run.font.name:
                        font_name = run.font.name
                    if run.font.size:
                        font_size = run.font.size.pt or 14.0

                lines.append(DocumentLine(text=text, font=font_name, size=font_size, page=1))

            # Парсинг таблиц внутри Word-файла
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            t = p.text.strip()
                            if t:
                                lines.append(DocumentLine(text=t, font="Times New Roman", size=12.0, page=1))
        except Exception as e:
            print(f"⚠️ Ошибка парсинга DOCX: {e}")

    # 2. Парсинг файлов PDF (.pdf)
    elif file_ext == 'pdf' and pypdf:
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page_idx, page in enumerate(reader.pages):
                text_content = page.extract_text()
                if not text_content:
                    continue
                for block in text_content.split('\n'):
                    if block.strip():
                        lines.append(
                            DocumentLine(text=block.strip(), font="Times New Roman", size=12.0, page=page_idx + 1))
        except Exception as e:
            print(f"⚠️ Ошибка парсинга PDF: {e}")

    # Запасной текстовый разбор, если библиотеки не справились или формат иной
    if not lines:
        try:
            text_content = file_bytes.decode('utf-8', errors='ignore')
            for chunk in text_content.split('\n'):
                if chunk.strip():
                    lines.append(DocumentLine(text=chunk.strip(), font="Arial", size=12.0, page=1))
        except Exception:
            pass

    # Железнобетонный дефолт, чтобы алгоритм не упал из-за пустого массива
    if not lines:
        lines.append(DocumentLine(text="Пустой документ ЛНА", font="Times New Roman", size=14.0, page=1))

    return lines


def analyze_document_text(doc_type: str, lines: List[DocumentLine]) -> Dict[str, Any]:
    """Анализирует текст на соответствие ГОСТ и НПА"""
    errors = []

    # Собираем весь текст в одну строку в нижнем регистре для поиска разделов
    text_all = " ".join([line.text for line in lines]).lower()

    # Умные регулярки (ловят инициалы, капс и текстовые даты)
    fio_pattern = re.compile(
        r"(?:[А-ЯЁ][а-яё\-]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.)|"
        r"(?:[А-ЯЁ]{3,}\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.)|"
        r"(?:[А-ЯЁ][а-яё\-]+\s+[А-ЯЁ][а-яё\-]+(?:\s+[А-ЯЁ][а-яё\-]+)?)"
    )
    date_pattern = re.compile(
        r"\b(0[1-9]|[12][0-9]|3[01])\.(0[1-9]|1[0-2])\.(\d{4})\b|"
        r"\b\d{4}\s*(?:г\.|года?)\b|"
        r"(?:0[1-9]|[12][0-9]|3[01])?\s*[а-яёА-ЯЁ]+\s*\d{4}"
    )

    has_fio = False
    has_date = False
    invalid_font_count = 0
    total_lines = len(lines) if lines else 1

    for line in lines:
        if fio_pattern.search(line.text):
            has_fio = True
        if date_pattern.search(line.text):
            has_date = True

        # Проверяем шрифты (допускаем Times New Roman и Arial)
        font_name = line.font.lower().replace(" ", "")
        if font_name and "times" not in font_name and "arial" not in font_name:
            invalid_font_count += 1

    if doc_type == "instruction":
        # 1. Проверка реквизитов
        if not has_fio:
            errors.append({"category": "Наличие ФИО", "severity": "CRITICAL",
                           "message": "В документе не обнаружены ФИО ответственных лиц", "location": "Шапка",
                           "fine_equivalent": "50 000 руб.", "legal_tip": "Укажите ФИО разработчиков."})
        if not has_date:
            errors.append({"category": "Дата", "severity": "WARNING", "message": "Не найдена дата утверждения",
                           "location": "Реквизиты", "fine_equivalent": "0 руб.",
                           "legal_tip": "Рекомендуется проставить дату."})
        if invalid_font_count > total_lines * 0.2:
            errors.append({"category": "Шрифт", "severity": "WARNING",
                           "message": "Более 20% текста использует нестандартный шрифт", "location": "Текст",
                           "fine_equivalent": "0 руб.",
                           "legal_tip": "Используйте Times New Roman 14pt (ГОСТ Р 7.0.97)."})

        # 2. Проверка 5 обязательных разделов (Приказ 772н)
        required_sections = [
            ("общие требования", "Общие требования охраны труда"),
            ("перед началом", "Требования перед началом работы"),
            ("во время", "Требования во время работы"),
            ("в аварийных", "Требования в аварийных ситуациях"),
            ("по окончании", "Требования по окончании работы")
        ]
        for pattern, name in required_sections:
            if pattern not in text_all:
                errors.append({
                    "category": "Структура (Приказ 772н)",
                    "severity": "CRITICAL",
                    "message": f"Отсутствует обязательный раздел: '{name}'",
                    "location": "Тело документа",
                    "fine_equivalent": "50 000 руб.",
                    "legal_tip": "ТК РФ требует наличия данного раздела в инструкции."
                })

    elif doc_type == "journal":
        if not has_fio:
            errors.append(
                {"category": "Наличие ФИО", "severity": "CRITICAL", "message": "Не обнаружены ФИО инструктируемых лиц",
                 "location": "Таблица", "fine_equivalent": "80 000 руб.",
                 "legal_tip": "Заполните обязательные графы журнала."})

    critical_count = sum(1 for e in errors if e["severity"] == "CRITICAL")
    warning_count = sum(1 for e in errors if e["severity"] == "WARNING")

    # Пересчет баллов
    base_score = 100
    if critical_count > 0:
        base_score -= (20 * critical_count)  # -20% за каждую критическую
    if warning_count > 0:
        base_score -= (10 * warning_count)  # -10% за ворнинги

    compliance_percent = max(0, min(100, base_score))

    return {
        "status": "FAILED" if critical_count > 0 else "PASSED",
        "total_errors": len(errors),
        "compliance_percent": compliance_percent,
        "critical_errors": critical_count,
        "warnings": warning_count,
        "errors_list": errors
    }


@router.post("/api/v1/verify/expert")
async def verify_document(
        doc_type: str = Form(...),
        file: UploadFile = File(...)
):
    try:
        file_bytes = await file.read()

        # 1. Извлекаем реальные строки из файла
        lines = extract_lines_from_file(file.filename, file_bytes)

        # 2. Запускаем аудит по ГОСТам и НПА
        analysis = analyze_document_text(doc_type, lines)

        # Вычисляем число пройденных критериев (условно)
        base_checks = 8 if doc_type == "instruction" else 3
        passed_count = max(0, base_checks - analysis["total_errors"])

        # 3. Формируем универсальный ответ (совместимый и с Pydantic, и с JS фронтенда)
        response_data = {
            "success": True,
            "status": analysis["status"],
            "filename": file.filename,
            "score": analysis["compliance_percent"],
            "compliance_percent": analysis["compliance_percent"],
            "critical_errors": analysis["critical_errors"],
            "warnings": analysis["warnings"],
            "passed_count": passed_count,
            "total_errors": analysis["total_errors"],
            "errors_list": analysis["errors_list"],
            "stats": {
                "critical_errors": analysis["critical_errors"],
                "warnings": analysis["warnings"],
                "verified_pages": len(set(line.page for line in lines))
            },
            "gosts": [
                {
                    "name": "ГОСТ Р 7.0.97 / Приказ 772н",
                    "status": "Соответствует" if analysis["critical_errors"] == 0 else "Нарушен",
                    "details": f"Выявлено {analysis['total_errors']} несоответствий."
                }
            ]
        }

        # Сохраняем в историю проверок для выгрузки Excel
        audit_history.append(response_data)
        return response_data

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Внутренняя ошибка парсера: {str(e)}")


@router.get("/api/v1/verify/export/{doc_type}")
async def export_audit_report(doc_type: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "Аудит ЛНА"

    # Стилизация Excel отчета
    header_fill = PatternFill(start_color="1A73E8", end_color="1A73E8", fill_type="solid")
    header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")

    headers = ["Файл", "Тип ЛНА", "Критерий проверки", "Статус", "Штрафной риск", "Рекомендация эксперта"]
    ws.append(headers)

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    row_num = 2
    # Берем данные из актуальной истории проверок
    for record in audit_history:
        if record.get("filename"):
            for err in record.get("errors_list", []):
                ws.append([
                    record["filename"],
                    "Инструкция" if doc_type == "instruction" else "Журнал",
                    err["category"],
                    err["severity"],
                    err["fine_equivalent"],
                    err["legal_tip"]
                ])
                for col_idx in range(1, 7):
                    ws.cell(row=row_num, column=col_idx).alignment = Alignment(vertical="top", wrap_text=True)
                row_num += 1

    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max(max_len + 3, 12), 40)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    raw_filename = f"OHS_Report_{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    encoded_filename = quote(raw_filename)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/api/v1/dashboard/director")
async def director_risk_dashboard() -> Dict[str, Any]:
    return {
        "total_documents": len(audit_history),
        "passed_documents": sum(1 for r in audit_history if r["status"] == "PASSED"),
        "failed_documents": sum(1 for r in audit_history if r["status"] == "FAILED"),
        "last_check": datetime.now().isoformat()
    }