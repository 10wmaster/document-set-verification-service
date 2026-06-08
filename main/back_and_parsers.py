import io
import re
import enum
import os
import tempfile
import shutil
from datetime import datetime, timedelta
from urllib.parse import quote
from typing import List, Optional, Dict, Any

from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.concurrency import run_in_threadpool
from pydantic import BaseModel

# Работа с Excel
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

# Парсеры
from docx import Document
import fitz  # PyMuPDF

app = FastAPI(
    title="OHS Expert Compliance Service API",
    description="MVP сервиса экспресс-аудита ЛНА по охране труда с учетом требований Приказа 772н, Постановления 2464 и ГОСТ 12.0.004-2015"
)


# === 1. ПЕРЕЧИСЛЕНИЯ ===
class SeverityLevel(str, enum.Enum):
    CRITICAL = "CRITICAL"
    WARNING = "WARNING"
    RECOMMENDATION = "RECOMMEND"


# === 2. СХЕМЫ ДАННЫХ (Pydantic-модели) ===
class ValidationErrorItem(BaseModel):
    category: str
    severity: SeverityLevel
    message: str
    location: Optional[str] = None
    fine_equivalent: Optional[str] = "0 руб."
    legal_tip: Optional[str] = None


class DocumentChunk(BaseModel):
    text: str
    font: str
    size: float
    page: Optional[int] = None
    is_table: bool = False  # Флаг принадлежности текста к таблице


class ValidationRequest(BaseModel):
    doc_type: str
    document_data: List[DocumentChunk]


# === 3. ПАРСЕРЫ ФАЙЛОВ ===
def parser_docx(file_path: str) -> List[dict]:
    doc = Document(file_path)
    result = []

    # 1. Чтение обычного текста (абзацев)
    for para in doc.paragraphs:
        p_fmt = para.paragraph_format
        para_data = {
            "alignment": str(p_fmt.alignment),
            "space_before": p_fmt.space_before.pt if p_fmt.space_before else 0.0,
            "space_after": p_fmt.space_after.pt if p_fmt.space_after else 0.0,
            "first_line": p_fmt.first_line_indent.cm if p_fmt.first_line_indent else 0.0,
            "line_spacing": p_fmt.line_spacing if p_fmt.line_spacing else 1.0
        }
        for run in para.runs:
            text = run.text.strip()
            if not text: continue
            result.append({
                "text": text,
                "font": run.font.name or 'Times New Roman',
                "size": run.font.size.pt if run.font.size else 12.0,
                "is_table": False,
                **para_data
            })

    # 2. Изолированное чтение таблиц (склеивание ранов ячейки в одну строку)
    processed_cells = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell in processed_cells:
                    continue
                processed_cells.add(cell)

                text = cell.text.strip().replace('\n', ' ')
                if not text: continue

                result.append({
                    "text": text,
                    "font": "Times New Roman",
                    "size": 12.0,
                    "is_table": True
                })
    return result


def parser_pdf(file_path: str) -> List[dict]:
    doc = fitz.open(file_path)
    result = []
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if b["type"] == 0:
                for line in b["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        if not text: continue
                        result.append({
                            "text": text,
                            "font": span["font"],
                            "size": round(span["size"], 1),
                            "is_bold": bool(span["flags"] & 2 ** 4),
                            "origin": span["origin"],
                            "page": page.number + 1,
                            "page_width_mm": round(page.rect.width * 25.4 / 72, 1),
                            "page_height_mm": round(page.rect.height * 25.4 / 72, 1),
                            "is_table": False
                        })
    doc.close()
    return result


# === 4. БИЗНЕС-ЛОГИКА (ПРОВЕРКИ С ГОСТ И ПОСТ. 2464) ===
def validate_gost_7_0_97(chunks: List[DocumentChunk]) -> List[ValidationErrorItem]:
    errors = []
    for idx, chunk in enumerate(chunks):
        if len(chunk.text.strip()) < 3 or chunk.is_table:
            continue
        font_lower = chunk.font.lower()
        if "times" not in font_lower and "arial" not in font_lower:
            errors.append(ValidationErrorItem(
                category="Нормоконтроль (ГОСТ 7.0.97)",
                severity=SeverityLevel.WARNING,
                message=f"Нестандартный шрифт '{chunk.font}'.",
                location=f"Стр. {chunk.page if chunk.page else idx + 1}",
                fine_equivalent="0 руб. (Прямой штраф отсутствует)",
                legal_tip="Для соответствия ГОСТ Р 7.0.97-2016 измените шрифт на Times New Roman or Arial."
            ))
    return errors


def validate_order_772n(chunks: List[DocumentChunk]) -> List[ValidationErrorItem]:
    errors = []
    full_text = " ".join([c.text.lower() for c in chunks])
    required_chapters = {
        "общие требования": "Общие требования охраны труда",
        "перед началом работы": "Требования охраны труда перед началом работы",
        "во время работы": "Требования охраны труда во время работы",
        "в аварийных ситуациях": "Требования охраны труда в аварийных ситуациях",
        "по окончании работы": "Требования охраны труда по окончании работы"
    }
    for key, original_name in required_chapters.items():
        if key not in full_text:
            errors.append(ValidationErrorItem(
                category="Структура (Приказ 772н)",
                severity=SeverityLevel.CRITICAL,
                message=f"Отсутствует обязательный нормативный раздел: '{original_name}'",
                location="Structures",
                fine_equivalent="от 50 000 до 80 000 руб. (ч. 1 ст. 5.27.1 КоАП РФ)",
                legal_tip=f"Добавьте в документ главу с точным заголовком: '{original_name}'."
            ))
    if "сиз" not in full_text and "средств индивидуальной защиты" not in full_text:
        errors.append(ValidationErrorItem(
            category="Оптимизация формулировок",
            severity=SeverityLevel.RECOMMENDATION,
            message="В тексте общих требований не конкретизирован порядок выдачи и применения СИЗ.",
            location="Раздел: Общие требования",
            fine_equivalent="0 руб. (Рекомендация)",
            legal_tip="Добавьте формулировку об обязанности правильно применять выданные СИЗ."
        ))
    if "первой помощ" not in full_text and "пострадавш" not in full_text:
        errors.append(ValidationErrorItem(
            category="Качество контента",
            severity=SeverityLevel.RECOMMENDATION,
            message="Раздел аварийных ситуаций содержит слабые формулировки по оказанию первой помощи.",
            location="Раздел: В аварийных ситуациях",
            fine_equivalent="0 руб. (Рекомендация)",
            legal_tip="Рекомендуется внедрить четкий алгоритм самопомощи при обнаружении пострадавшего."
        ))
    return errors


def validate_journal_post_2464(chunks: List[DocumentChunk]) -> List[ValidationErrorItem]:
    errors = []
    full_text = " ".join([c.text.lower() for c in chunks])
    has_tables = any(c.is_table for c in chunks)

    # -------------------------------------------------------------
    # МАРКЕР 1: Проверка реквизитов (Постановление 2464 + Синонимы)
    # -------------------------------------------------------------
    if not any(kw in full_text for kw in ["фио", "фамилия", "имя", "отчество"]):
        errors.append(ValidationErrorItem(
            category="Реквизиты документа",
            severity=SeverityLevel.CRITICAL,
            message="В форме фиксации инструктажа не найден обязательный столбец: 'ФИО' сотрудника.",
            location="Шапка таблицы",
            fine_equivalent="от 110 000 до 130 000 руб. (ч. 3 ст. 5.27.1 КоАП РФ)",
            legal_tip="Без поля 'ФИО' инспекция ГИТ признает проведенный инструктаж недействительным."
        ))

    if not any(kw in full_text for kw in ["дата", "число"]):
        errors.append(ValidationErrorItem(
            category="Реквизиты документа",
            severity=SeverityLevel.CRITICAL,
            message="В форме фиксации инструктажа не найден обязательный столбец: 'ДАТА'.",
            location="Шапка таблицы",
            fine_equivalent="от 110 000 до 130 000 руб. (ч. 3 ст. 5.27.1 КоАП РФ)",
            legal_tip="Добавьте колонку для фиксации точной даты проведения инструктажа."
        ))

    # -------------------------------------------------------------
    # МАРКЕР 2: Новые валидаторы, обоснованные ГОСТ 12.0.004-2015
    # -------------------------------------------------------------
    # Глубокий анализ разделения подписей (ГОСТ 12.0.004-2015 + п. 87 Пост. 2464)
    has_signature = any(kw in full_text for kw in ["подпись", "распись", "расписался"])
    if has_signature:
        # ГОСТ требует фиксации подписей как инструктируемого, так и инструктирующего лица отдельно
        instruktiruemyi = any(kw in full_text for kw in ["инструктируемого", "работника", "обучаемого"])
        instruktiruyushiy = any(
            kw in full_text for kw in ["инструктирующего", "руководителя", "проверяющего", "мастера"])

        if not (instruktiruemyi and instruktiruyushiy):
            errors.append(ValidationErrorItem(
                category="Нормоконтроль (ГОСТ 12.0.004-2015)",
                severity=SeverityLevel.WARNING,
                message="Не разделены зоны ответственности подписей (Инструктируемого и Инструктирующего).",
                location="Табличная часть (Подписи)",
                fine_equivalent="0 руб. (Высокий регуляторный риск)",
                legal_tip="Согласно Приложению А ГОСТ 12.0.004-2015, в таблице должны быть явно выделены две разные графы для подписей."
            ))
    else:
        errors.append(ValidationErrorItem(
            category="Реквизиты документа",
            severity=SeverityLevel.CRITICAL,
            message="В форме фиксации инструктажа не найден обязательный столбец: 'ПОДПИСЬ'.",
            location="Шапка таблицы",
            fine_equivalent="от 110 000 до 130 000 руб. (ч. 3 ст. 5.27.1 КоАП РФ)",
            legal_tip="Подпись сотрудника — единственное юридическое доказательство факта проведения обучения."
        ))

    # Аудит фиксации оснований для внеплановых мероприятий (Требование ГОСТ)
    if "внеплан" in full_text:
        if not any(kw in full_text for kw in ["причина", "основание", "номер приказа", "№ приказа"]):
            errors.append(ValidationErrorItem(
                category="Нормоконтроль (ГОСТ 12.0.004-2015)",
                severity=SeverityLevel.RECOMMENDATION,
                message="В журнале отсутствует графа 'Причина проведения внепланового инструктажа'.",
                location="Шапка таблицы",
                fine_equivalent="0 руб. (Рекомендация)",
                legal_tip="ГОСТ 12.0.004-2015 рекомендует фиксировать номер приказа Минтруда или локального акта, послужившего триггером внепланового аудита."
            ))

    # -------------------------------------------------------------
    # МАРКЕР 3: Умный поиск дат (Защита от дат стандартов)
    # -------------------------------------------------------------
    date_pattern = r"\b\d{2}\.\d{2}\.\d{4}\b"
    current_year = datetime.now().year

    for chunk in chunks:
        # ЗАЩИТА 1: Если в docx есть таблицы, ищем даты ТОЛЬКО внутри ячеек таблиц!
        if has_tables and not chunk.is_table:
            continue

        matches = re.findall(date_pattern, chunk.text)
        for date_str in matches:
            try:
                date_obj = datetime.strptime(date_str, "%d.%m.%Y")

                # ЗАЩИТА 2: Отсекаем старые архивные нормативные даты (ГОСТ 2013 г., приказы 2020 г.)
                # Если дата старше чем текущая дата минус 2 года (730 дней) — это не запись сотрудника, игнорируем её.
                if date_obj < datetime.now() - timedelta(days=730):
                    continue

                # Если дата "свежая", но ей больше 6 месяцев — это реальная просрочка периодичности
                if datetime.now() - date_obj > timedelta(days=180):
                    errors.append(ValidationErrorItem(
                        category="Сроки действия (Пост. 2464)",
                        severity=SeverityLevel.CRITICAL,
                        message=f"Нарушена периодичность обучения! Инструктаж от {date_str} просрочен.",
                        location=f"Строка таблицы: '{chunk.text[:35]}...'",
                        fine_equivalent="от 110 000 до 130 000 руб. ЗА КАЖДОГО сотрудника (ч. 3 ст. 5.27.1 КоАП РФ)",
                        legal_tip="Немедленно отстраните работников с просроченными датами от выполнения обязанностей и проведите повторный инструктаж."
                    ))
            except ValueError:
                continue

    if not any(kw in full_text for kw in ["номер", "№", "п/п"]):
        errors.append(ValidationErrorItem(
            category="Structure (Постановление 2464)",
            severity=SeverityLevel.RECOMMENDATION,
            message="В таблице отсутствует сквозная нумерация записей (№ п/п).",
            location="Табличная часть",
            fine_equivalent="0 руб.",
            legal_tip="Добавьте первую колонку: '№ п/п' (Порядковый номер записи) для исключения фальсификации записей задним числом."
        ))
    return errors


# === 5. СИНХРОННАЯ ОБЕРТКА ДЛЯ ПАРСИНГА И ПРОВЕРКИ ===
def process_document_sync(file_path: str, ext: str, doc_type: str) -> Dict[str, Any]:
    document_data = []

    if ext == ".docx":
        raw_data = parser_docx(file_path)
        for c in raw_data:
            document_data.append(DocumentChunk(
                text=c["text"],
                font=c["font"],
                size=float(c["size"]),
                page=1,
                is_table=c.get("is_table", False)
            ))
    elif ext == ".pdf":
        raw_data = parser_pdf(file_path)
        for c in raw_data:
            document_data.append(DocumentChunk(
                text=c["text"],
                font=c["font"],
                size=float(c["size"]),
                page=c.get("page", 1),
                is_table=False
            ))

    all_violations = []
    if doc_type == "instruction":
        all_violations.extend(validate_gost_7_0_97(document_data))
        all_violations.extend(validate_order_772n(document_data))
    elif doc_type == "journal":
        all_violations.extend(validate_journal_post_2464(document_data))

    critical_count = sum(1 for v in all_violations if v.severity == SeverityLevel.CRITICAL)

    return {
        "status": "FAILED" if critical_count > 0 else "PASSED",
        "analyzed_at": datetime.now().isoformat(),
        "total_errors": len(all_violations),
        "violations": [v.dict() for v in all_violations]
    }


# === 6. ЭНДПОИНТЫ API ===
@app.post("/api/v1/verify/upload", summary="Загрузить и проверить файл (.docx, .pdf)")
async def verify_uploaded_file(
        file: UploadFile = File(...),
        doc_type: str = Form(..., description="Тип документа: 'instruction' или 'journal'")
):
    if doc_type not in ["instruction", "journal"]:
        raise HTTPException(status_code=400, detail="Тип документа должен быть 'instruction' или 'journal'")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".docx", ".pdf"]:
        raise HTTPException(status_code=400, detail="Поддерживаются только форматы .docx и .pdf")

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
        while content := await file.read(1024 * 1024):
            temp_file.write(content)
        temp_file_path = temp_file.name

    try:
        result = await run_in_threadpool(
            process_document_sync,
            temp_file_path,
            ext,
            doc_type
        )
        result["filename"] = file.filename
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка при обработке файла: {str(e)}")
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)


@app.post("/api/v1/verify/expert", response_model=Dict[str, Any], summary="Проверить готовый массив чанков (JSON)")
async def expert_document_audit(request: ValidationRequest):
    all_violations = []
    if request.doc_type == "instruction":
        all_violations.extend(validate_gost_7_0_97(request.document_data))
        all_violations.extend(validate_order_772n(request.document_data))
    elif request.doc_type == "journal":
        all_violations.extend(validate_journal_post_2464(request.document_data))
    else:
        raise HTTPException(status_code=400, detail="Указан неподдерживаемый тип документа")

    critical_count = sum(1 for v in all_violations if v.severity == SeverityLevel.CRITICAL)

    return {
        "status": "FAILED" if critical_count > 0 else "PASSED",
        "analyzed_at": datetime.now().isoformat(),
        "total_errors": len(all_violations),
        "violations": all_violations
    }


@app.post("/api/v1/verify/export", summary="Экспорт отчета в Excel")
async def export_compliance_report(request: ValidationRequest):
    if request.doc_type not in ["instruction", "journal"]:
        raise HTTPException(status_code=400, detail="Неверный тип документа для экспорта")

    if request.doc_type == "instruction":
        violations = validate_order_772n(request.document_data) + validate_gost_7_0_97(request.document_data)
    else:
        violations = validate_journal_post_2464(request.document_data)

    wb = Workbook()
    ws = wb.active
    ws.title = "Протокол комплаенса"
    ws.sheet_view.showGridLines = True

    headers = ["Категория", "Уровень риска", "Описание нарушения", "Локация", "Штраф по КоАП РФ",
               "Рекомендация по исправлению"]
    ws.append(headers)

    header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    row_num = 2
    for v in violations:
        ws.append([
            v.category,
            v.severity.value,
            v.message,
            v.location or "Не указано",
            v.fine_equivalent,
            v.legal_tip
        ])
        severity_cell = ws.cell(row=row_num, column=2)
        if v.severity == SeverityLevel.CRITICAL:
            severity_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            severity_cell.font = Font(color="9C0006", bold=True)
        elif v.severity == SeverityLevel.WARNING:
            severity_cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
            severity_cell.font = Font(color="9C6500", bold=True)
        else:
            severity_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            severity_cell.font = Font(color="006100", bold=True)

        for col_num in range(1, 7):
            ws.cell(row=row_num, column=col_num).alignment = Alignment(vertical="top", wrap_text=True)
        row_num += 1

    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = min(max(max_len + 3, 12), 40)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    raw_filename = f"OHS_Report_{request.doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    encoded_filename = quote(raw_filename)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@app.get("/api/v1/dashboard/director", summary="Сводка для дашборда")
async def director_risk_dashboard():
    return {
        "company_status": "HIGH_RISK",
        "compliance_score_percent": 68,
        "git_inspection_readiness": "Не готов",
        "financial_risk_estimation": "До 1 300 000 рублей (на основе ст. 5.27.1 КоАП за дефекты в журналах)",
        "recommendation": "Срочно исправьте критические нарушения: добавьте недостающие разделы в инструкциях по Приказу 772н и обновите даты просроченных инструктажей в журналах."
    }