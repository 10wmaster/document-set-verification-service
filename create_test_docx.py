from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_fake_kursovaya(filename):
    doc = Document()

    # 1. Шапка
    p = doc.add_paragraph("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n" * 3) # Отступы

    # 2. Тема
    p_type = doc.add_paragraph("КУРСОВАЯ РАБОТА")
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_topic = doc.add_paragraph("Тема: Автоматизация проверки документов на Python")
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n" * 2)

    # 3. Блок с данными (создаем таблицу 2х2 без границ)
    table = doc.add_table(rows=2, cols=2)
    
    # Левая колонка — подписи, правая — данные
    table.cell(0, 0).text = "Выполнил студент:"
    table.cell(0, 1).text = "Иванов Иван Иванович" # Твое целевое ФИО
    
    table.cell(1, 0).text = "Группа:"
    table.cell(1, 1).text = "ИБ-123"

    doc.save(filename)
    print(f"Файл {filename} создан!")

create_fake_kursovaya("test_titul.docx")